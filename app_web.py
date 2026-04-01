import streamlit as st
import pdfplumber
import re
import io
import os
import time
from docx import Document
from docx.shared import Pt, RGBColor
from openai import OpenAI

# =========================
# CONFIGURAÇÃO DE AMBIENTE
# =========================
api_key = st.secrets.get("OPENAI_API_KEY") if "OPENAI_API_KEY" in st.secrets else os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key) if api_key else None

# =========================
# FUNÇÕES DE PROCESSAMENTO
# =========================

def extrair_texto_pdf(pdf):
    texto = ""
    with pdfplumber.open(pdf) as p:
        for page in p.pages:
            texto += (page.extract_text() or "") + "\n"
    return texto

def dividir_texto(texto, max_chars=8000):
    """Divide textos grandes para a IA ler em pedaços seguros."""
    blocos = []
    inicio = 0
    while inicio < len(texto):
        fim = inicio + max_chars
        if fim < len(texto):
            pos_quebra = texto.rfind('\n\n', inicio, fim)
            if pos_quebra == -1 or pos_quebra <= inicio:
                pos_quebra = texto.rfind('\n', inicio, fim)
            if pos_quebra != -1 and pos_quebra > inicio:
                fim = pos_quebra
        blocos.append(texto[inicio:fim])
        inicio = fim
    return blocos

def gerar_word_fidelidade_total(texto_final):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.0 

    for linha in texto_final.split('\n'):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
        
        # Preserva recuos
        if linha.startswith("Art."):
            p.paragraph_format.first_line_indent = Pt(0)
        elif linha.startswith("§") or re.match(r'^[a-z]\)|\d+\.|[IVXLC]+\s?-', linha):
            p.paragraph_format.left_indent = Pt(36)

        partes = re.split(r'(~~.*?~~|\*\*.*?\*\*|\[\[.*?\]\])', linha)
        
        for parte in partes:
            if parte.startswith('~~') and parte.endswith('~~'):
                run = p.add_run(parte[2:-2])
                run.font.strike = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif parte.startswith('**') and parte.endswith('**'):
                run = p.add_run(parte[2:-2])
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            elif parte.startswith('[[') and parte.endswith(']]'):
                run = p.add_run(parte[2:-2])
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                run = p.add_run(parte)
                run.font.color.rgb = RGBColor(0, 0, 0)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def processar_comparacao_rapida(texto_base, texto_alteracoes):
    blocos = dividir_texto(texto_base, 8000)
    resultado_final = []
    
    total_blocos = len(blocos)
    barra_progresso = st.progress(0)
    status_texto = st.empty()

    # O Segredo da Velocidade: Se não houver alteração no bloco, a IA devolve 1 palavra só.
    prompt_sistema = """
    Você é um compilador jurídico veloz e preciso. 
    Sua missão é aplicar o 'TEXTO 2 (Alterações)' no 'BLOCO DO TEXTO 1'.

    REGRA DE VELOCIDADE (CRÍTICA):
    Avalie o bloco. Se o TEXTO 2 NÃO determinar NENHUMA alteração, revogação ou inclusão para os artigos contidos exatamente neste bloco, VOCÊ DEVE RESPONDER APENAS A PALAVRA:
    NENHUMA_ALTERACAO
    
    REGRA DE CONSOLIDAÇÃO (Caso haja mudança no bloco):
    1. Reescreva o bloco inteiro.
    2. O que sai fica riscado: ~~texto antigo~~ (Revogado pela [[Nome da Portaria]]).
    3. O que entra fica em negrito: **texto novo** (Incluído pela [[Nome da Portaria]]).
    4. Nomes de Portarias alteradoras SEMPRE em colchetes duplos [[ ]].
    5. Mantenha os artigos que não sofreram mudança iguais.
    """

    for i, bloco in enumerate(blocos):
        status_texto.markdown(f"**Escaneando e processando bloco {i+1} de {total_blocos}...**")
        
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": prompt_sistema},
                    {"role": "user", "content": f"BLOCO DO TEXTO 1:\n{bloco}\n\nTEXTO 2 (ALTERAÇÕES):\n{texto_alteracoes[:15000]}"}
                ],
                temperature=0
            )
            
            resposta_ia = response.choices[0].message.content.strip()
            
            # Se a IA não identificou mudança, devolvemos o bloco original intacto e rápido!
            if "NENHUMA_ALTERACAO" in resposta_ia.upper():
                resultado_final.append(bloco)
            else:
                resultado_final.append(resposta_ia)
                
        except Exception as e:
            resultado_final.append(bloco) # Em caso de erro de rede, salva o original para não perder o texto
            st.toast(f"Pequeno erro de rede no bloco {i+1}, original mantido.")
            
        barra_progresso.progress((i + 1) / total_blocos)
        time.sleep(0.5)

    status_texto.empty()
    return "\n\n".join(resultado_final)

# =========================
# INTERFACE STREAMLIT
# =========================

st.set_page_config(page_title="Consolidador Ultra-Rápido", layout="wide")
st.title("⚖️ Consolidador Jurídico de Alta Velocidade")
st.info("Otimizado para portarias gigantes. Pula automaticamente as páginas que não sofreram alterações.")

if 'resultado_docx' not in st.session_state:
    st.session_state.resultado_docx = None
if 'texto_tela' not in st.session_state:
    st.session_state.texto_tela = None

col1, col2 = st.columns(2)
with col1:
    pdf_base = st.file_uploader("1. Portaria ORIGINAL (Ex: 100+ páginas)", type="pdf", key="f1")
with col2:
    pdf_alt = st.file_uploader("2. Portaria ALTERADORA", type="pdf", key="f2")

if st.button("🚀 Processar Rapidamente", key="btn_run"):
    if not pdf_base or not pdf_alt:
        st.warning("Faça o upload dos dois arquivos PDF.")
    elif not client:
        st.error("API Key da OpenAI não configurada.")
    else:
        st.session_state.resultado_docx = None
        
        with st.spinner("Extraindo textos..."):
            t_base = extrair_texto_pdf(pdf_base)
            t_alt = extrair_texto_pdf(pdf_alt)
        
        # Executa a nova lógica rápida
        texto_processado = processar_comparacao_rapida(t_base, t_alt)
        st.session_state.texto_tela = texto_processado
        
        with st.spinner("Gerando arquivo Word perfeitamente formatado..."):
            st.session_state.resultado_docx = gerar_word_fidelidade_total(texto_processado)

# Área de Download segura
if st.session_state.resultado_docx:
    st.success("✅ Consolidação concluída e arquivo Word gerado!")
    
    st.download_button(
        label="📥 Baixar Portaria_Consolidada.docx",
        data=st.session_state.resultado_docx,
        file_name="Portaria_Consolidada_Formatada.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="btn_down"
    )
    
    with st.expander("Ver prévia rápida do texto na tela"):
        st.write(st.session_state.texto_tela)